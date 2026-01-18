"""
Tiptap JSON 格式轉換工具
支援將 Tiptap JSON 格式轉換為 Word 文檔格式（透過 python-docx）
"""
from typing import Dict, Any, List, Optional
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Tiptap conversion will be disabled.")


class TiptapConverter:
    """
    Tiptap JSON 格式轉換器
    支援加粗文字、巢狀列表等複雜格式
    """
    
    @staticmethod
    def json_to_word_content(tiptap_json: Dict[str, Any]) -> str:
        """
        將 Tiptap JSON 格式轉換為純文本（用於模板渲染）
        
        Args:
            tiptap_json: Tiptap JSON 格式的字典
            
        Returns:
            純文本字符串
        """
        if not tiptap_json or "content" not in tiptap_json:
            return ""
        
        return TiptapConverter._parse_content(tiptap_json.get("content", []))
    
    @staticmethod
    def _parse_content(content: List[Dict[str, Any]]) -> str:
        """遞迴解析 Tiptap content 節點"""
        result = []
        
        for node in content:
            node_type = node.get("type", "")
            
            if node_type == "paragraph":
                text = TiptapConverter._parse_paragraph(node)
                if text:
                    result.append(text)
            
            elif node_type == "heading":
                level = node.get("attrs", {}).get("level", 1)
                text = TiptapConverter._parse_paragraph(node)
                if text:
                    # 根據標題級別添加前綴
                    prefix = "#" * level + " "
                    result.append(prefix + text)
            
            elif node_type in ["bulletList", "orderedList"]:
                items = TiptapConverter._parse_list(node, node_type == "orderedList")
                result.extend(items)
            
            elif node_type == "codeBlock":
                text = TiptapConverter._extract_text_from_node(node)
                if text:
                    result.append(f"```\n{text}\n```")
            
            elif node_type == "blockquote":
                text = TiptapConverter._parse_content(node.get("content", []))
                if text:
                    result.append(f"> {text}")
            
            elif node_type == "hardBreak":
                result.append("\n")
            
            else:
                # 處理其他節點類型，提取文本
                text = TiptapConverter._extract_text_from_node(node)
                if text:
                    result.append(text)
        
        return "\n\n".join(result)
    
    @staticmethod
    def _parse_paragraph(node: Dict[str, Any]) -> str:
        """解析段落節點，處理加粗等格式化"""
        return TiptapConverter._extract_text_from_node(node)
    
    @staticmethod
    def _extract_text_from_node(node: Dict[str, Any]) -> str:
        """從節點中提取文本，處理加粗、斜體等格式"""
        content = node.get("content", [])
        if not content:
            # 如果是文本節點
            if node.get("type") == "text":
                text = node.get("text", "")
                marks = node.get("marks", [])
                # 應用格式標記（加粗、斜體等）
                for mark in marks:
                    mark_type = mark.get("type", "")
                    if mark_type == "bold":
                        text = f"**{text}**"
                    elif mark_type == "italic":
                        text = f"*{text}*"
                    elif mark_type == "code":
                        text = f"`{text}`"
                    elif mark_type == "underline":
                        text = f"__{text}__"
                    elif mark_type == "strike":
                        text = f"~~{text}~~"
                return text
            return ""
        
        # 遞迴處理子節點
        parts = []
        for child in content:
            parts.append(TiptapConverter._extract_text_from_node(child))
        
        return "".join(parts)
    
    @staticmethod
    def _parse_list(node: Dict[str, Any], ordered: bool = False, level: int = 0) -> List[str]:
        """
        解析列表節點，支援巢狀列表
        
        Args:
            node: 列表節點
            ordered: 是否為有序列表
            level: 巢狀層級（用於縮進）
            
        Returns:
            格式化後的列表項列表
        """
        items = []
        content = node.get("content", [])
        indent = "  " * level  # 每層縮進 2 個空格
        
        item_counter = 1
        
        for list_item in content:
            if list_item.get("type") == "listItem":
                item_content = list_item.get("content", [])
                item_texts = []
                
                for item_node in item_content:
                    if item_node.get("type") == "paragraph":
                        text = TiptapConverter._extract_text_from_node(item_node)
                        if text:
                            item_texts.append(text)
                    elif item_node.get("type") in ["bulletList", "orderedList"]:
                        # 處理巢狀列表
                        nested_items = TiptapConverter._parse_list(
                            item_node, 
                            item_node.get("type") == "orderedList",
                            level + 1
                        )
                        item_texts.extend(nested_items)
                
                if item_texts:
                    # 構建列表項前綴
                    if ordered:
                        prefix = f"{item_counter}."
                        item_counter += 1
                    else:
                        prefix = "-"
                    
                    # 第一行有前綴，後續行對齊
                    first_line = f"{indent}{prefix} {item_texts[0]}"
                    items.append(first_line)
                    
                    # 處理多行內容
                    for line in item_texts[1:]:
                        # 多行內容需要適當縮進
                        spaces = len(prefix) + 1
                        items.append(f"{indent}{' ' * spaces}{line}")
        
        return items
    
    @staticmethod
    def json_to_word_doc(tiptap_json: Dict[str, Any], output_path: str) -> None:
        """
        將 Tiptap JSON 格式直接轉換為 Word 文檔（使用 python-docx）
        支援加粗文字和巢狀列表的格式保留
        
        Args:
            tiptap_json: Tiptap JSON 格式的字典
            output_path: 輸出 Word 文檔路徑
            
        Returns:
            None（直接保存文件）
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available. Cannot create Word document.")
        
        doc = Document()
        content = tiptap_json.get("content", [])
        
        TiptapConverter._add_content_to_doc(doc, content)
        
        doc.save(output_path)
    
    @staticmethod
    def _add_content_to_doc(doc: Document, content: List[Dict[str, Any]]) -> None:
        """將 Tiptap content 添加到 Word 文檔"""
        for node in content:
            node_type = node.get("type", "")
            
            if node_type == "paragraph":
                para = doc.add_paragraph()
                TiptapConverter._add_paragraph_content(para, node)
            
            elif node_type == "heading":
                level = node.get("attrs", {}).get("level", 1)
                # Word 文檔有 9 個標題級別
                heading_style = f"Heading {min(level, 9)}"
                para = doc.add_paragraph(style=heading_style)
                TiptapConverter._add_paragraph_content(para, node)
            
            elif node_type in ["bulletList", "orderedList"]:
                TiptapConverter._add_list_to_doc(doc, node, node_type == "orderedList")
            
            elif node_type == "codeBlock":
                para = doc.add_paragraph(style="Intense Quote")
                text = TiptapConverter._extract_text_from_node(node)
                para.add_run(text)
            
            elif node_type == "blockquote":
                para = doc.add_paragraph(style="Intense Quote")
                TiptapConverter._add_paragraph_content(para, node)
            
            elif node_type == "hardBreak":
                doc.add_paragraph()
    
    @staticmethod
    def _add_paragraph_content(para, node: Dict[str, Any]) -> None:
        """向段落添加內容，處理加粗等格式"""
        content = node.get("content", [])
        for child in content:
            TiptapConverter._add_text_with_marks(para, child)
    
    @staticmethod
    def _add_text_with_marks(para, node: Dict[str, Any]) -> None:
        """添加帶格式標記的文本"""
        if node.get("type") == "text":
            text = node.get("text", "")
            marks = node.get("marks", [])
            
            run = para.add_run(text)
            
            # 應用格式標記
            for mark in marks:
                mark_type = mark.get("type", "")
                if mark_type == "bold":
                    run.bold = True
                elif mark_type == "italic":
                    run.italic = True
                elif mark_type == "underline":
                    run.underline = True
                elif mark_type == "code":
                    run.font.name = "Courier New"
            
        elif "content" in node:
            for child in node.get("content", []):
                TiptapConverter._add_text_with_marks(para, child)
    
    @staticmethod
    def _add_list_to_doc(doc: Document, node: Dict[str, Any], ordered: bool = False) -> None:
        """向文檔添加列表，支援巢狀"""
        content = node.get("content", [])
        
        for list_item in content:
            if list_item.get("type") == "listItem":
                item_content = list_item.get("content", [])
                
                for item_node in item_content:
                    if item_node.get("type") == "paragraph":
                        if ordered:
                            para = doc.add_paragraph(style="List Number")
                        else:
                            para = doc.add_paragraph(style="List Bullet")
                        TiptapConverter._add_paragraph_content(para, item_node)
                    
                    elif item_node.get("type") in ["bulletList", "orderedList"]:
                        # 巢狀列表：在 Word 中，可以通過調整列表級別實現
                        # 這裡簡化處理，直接遞迴添加
                        TiptapConverter._add_list_to_doc(
                            doc, 
                            item_node, 
                            item_node.get("type") == "orderedList"
                        )
